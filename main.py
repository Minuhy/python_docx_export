# 执行 pip install -r requirements.txt 安装依赖
import os
import docx
from oletools import oleobj


def re_decode(s: str, encoding: str = 'gbk'):
    """
    重新解码，解决oleobj对中文乱码的问题
    :param s: 原始字符串
    :param encoding: 新的解码编码，默认为 GBK
    :return: 新的字符串
    """
    i81 = s.encode('iso-8859-1')
    return i81.decode(encoding)


def get_new_path(i: int, file_path: str, out_dir: str = __file__, file_name_format=None):
    """
    根据序号和文件名拿到新的路径名，
    默认输出目录为当前文件所在文件夹
    :param i: 序号
    :param file_path: 原始文件名（路径）
    :param out_dir: 输出目录，默认为当前文件所在文件夹
    :return: 新文件名（格式为“序号 - 原文件名”）
    :param file_name_format: 导出文件名格式，会直接替换对应的文字
    """
    if file_name_format is None:
        file_name_format = '|自增编号||连接符||原文件名||后缀名|'
    out_dir = os.path.dirname(os.path.abspath(out_dir))  # 获取导出路径
    number = str(i).rjust(4, '0')  # |自增编号|，不够前面添0
    link_char = ' - '  # |连接符|
    raw_file_name = os.path.basename(file_path)  # |原文件名|
    suffix = ''  # |后缀名|
    if '.' in raw_file_name:
        name_part = raw_file_name.split('.')
        raw_file_name = name_part[0]
        suffix = name_part[1]
    return out_dir + '\\' + file_name_format.replace('|自增编号|', number) \
        .replace('|连接符|', link_char) \
        .replace('|原文件名|', raw_file_name) \
        .replace('|后缀名|', '.' + suffix) \
        .strip()  # 拼接路径


def export_docx(docx_file: str, output_dir: str, name_format: str,
                export_type=None,
                is_del_raw=False,
                is_print=False
                ):
    """
    导出docx中的所有内容
    :param docx_file: docx路径
    :param output_dir: 导出的文件夹路径
    :param name_format: 导出文件名的格式
    :param export_type: 导出类型
    :param is_del_raw: 导出成功后是否删除原文件
    :param is_print: 是否打印调试信息
    :return: 导出的文件列表
    """

    # 文件导出列表
    if export_type is None:
        export_type = ['文本', '表格', '图片', '附件', '合并']
    export_files = []

    # 打开docx文档
    docx_document = docx.Document(docx_file)
    if is_print:
        print('打开文档完成')

    # 文档信息
    docx_properties = docx_document.core_properties
    all_properties = ''
    all_properties += '作者：' + str(docx_properties.author) + '\n'
    all_properties += '类别：' + str(docx_properties.category) + '\n'
    all_properties += '注释：' + str(docx_properties.comments) + '\n'
    all_properties += '内容状态：' + str(docx_properties.content_status) + '\n'
    all_properties += '创建时间：' + str(docx_properties.created) + '\n'
    all_properties += '标识符：' + str(docx_properties.identifier) + '\n'
    all_properties += '关键字：' + str(docx_properties.keywords) + '\n'
    all_properties += '语言：' + str(docx_properties.language) + '\n'
    all_properties += '最后修改者：' + str(docx_properties.last_modified_by) + '\n'
    all_properties += '上次打印：' + str(docx_properties.last_printed) + '\n'
    all_properties += '修改时间：' + str(docx_properties.modified) + '\n'
    all_properties += '修订：' + str(docx_properties.revision) + '\n'
    all_properties += '主题：' + str(docx_properties.subject) + '\n'
    all_properties += '标题：' + str(docx_properties.title) + '\n'
    all_properties += '版本：' + str(docx_properties.version) + '\n'
    if is_print:
        print('文档信息：', all_properties.replace('\n', '， '))

    # 导出文档信息
    info_file_path = get_new_path(len(export_files) + 1, 'docx文档信息.txt')
    with open(info_file_path, 'w', encoding='utf-8') as f:
        f.write(all_properties)
    export_files.append(info_file_path)

    # 所有文本
    all_text = ''
    for paragraph in docx_document.paragraphs:
        all_text += paragraph.text + ' '  # 段落之间用空格隔开
    if is_print:
        print('所有文本：', all_text)
    all_table_text = ''
    for table in docx_document.tables:
        for cell in getattr(table, '_cells'):
            all_table_text += cell.text + ' '  # 单元格之间用空格隔开
    if is_print:
        print('所有表格文本：', all_table_text)

    # 导出文本
    text_file_path = get_new_path(len(export_files) + 1, 'docx文本.txt')
    with open(text_file_path, 'w', encoding='utf-8') as f:
        f.write(all_text)
        f.write('\n')  # 换个行
        f.write(all_table_text)
    export_files.append(text_file_path)

    # 遍历所有附件
    docx_related_parts = docx_document.part.related_parts
    for part in docx_related_parts:
        part = docx_related_parts[part]
        part_name = str(part.partname)  # 附件路径（partname）
        if part_name.startswith('/word/media/') or part_name.startswith('/word/embeddings/'):  # 只导出这两个目录下的
            # 构建导出路径
            save_path = get_new_path(len(export_files) + 1, part.partname, output_dir)

            # ole 文件判断
            # 不符合 .bin 作为后缀且文件名中有ole，则不被认为是OLE文件
            if not (save_path.endswith('.bin') and 'ole' in save_path.lower()):
                # 直接写入文件
                if is_print:
                    print('DOCX 导出路径：', save_path)
                with open(save_path, 'wb') as f:
                    f.write(part.blob)
                export_files.append(save_path)  # 记录文件
            else:
                # 将字节数组传递给oleobj处理
                for ole in oleobj.find_ole(save_path, part.blob):
                    if ole is None:  # 没有找到 OLE 文件，跳过
                        continue

                    for path_parts in ole.listdir():  # 遍历OLE中的文件

                        # 判断是不是[1]Ole10Native，使用列表推导式忽略大小写，不是的话就不要继续了
                        if '\x01ole10native'.casefold() not in [path_part.casefold() for path_part in path_parts]:
                            continue

                        stream = None
                        try:
                            # 使用 Ole File 打开 OLE 文件
                            stream = ole.openstream(path_parts)
                            opkg = oleobj.OleNativeStream(stream)
                        except IOError:
                            if is_print:
                                print('不是OLE文件：', path_parts)
                            if stream is not None:  # 关闭文件流
                                stream.close()
                            continue

                        # 打印信息
                        if opkg.is_link:
                            if is_print:
                                print('是链接而不是文件，跳过')
                            continue

                        ole_filename = re_decode(opkg.filename)
                        ole_src_path = re_decode(opkg.src_path)
                        ole_temp_path = re_decode(opkg.temp_path)
                        if is_print:
                            print('文件名：{0}，原路径：{1}，缓存路径：{2}'.format(ole_filename, ole_src_path, ole_temp_path))

                        # 生成新的文件名
                        filename = get_new_path(len(export_files) + 1, ole_filename, output_dir)
                        if is_print:
                            print('OLE 导出路径：', filename)

                        # 转存
                        try:
                            if is_print:
                                print('导出OLE中的文件：', filename)
                            with open(filename, 'wb') as writer:
                                n_dumped = 0
                                next_size = min(oleobj.DUMP_CHUNK_SIZE, opkg.actual_size)
                                while next_size:
                                    data = stream.read(next_size)
                                    writer.write(data)
                                    n_dumped += len(data)
                                    if len(data) != next_size:
                                        if is_print:
                                            print('想要读取 {0}, 实际取得 {1}'.format(next_size, len(data)))
                                        break
                                    next_size = min(oleobj.DUMP_CHUNK_SIZE, opkg.actual_size - n_dumped)
                            export_files.append(filename)  # 记录导出的文件
                        except Exception as exc:
                            if is_print:
                                print('在转存时出现错误：{0} {1}'.format(filename, exc))
                        finally:
                            stream.close()
    if is_print:
        print('导出的所有文件：', export_files)
    return export_files


if __name__ == '__main__':
    file = 'docx/word.docx'
    files = export_docx(file, __file__, is_print=False)  # 打印调试信息（批量时会慢）
    print(files)
