# ! /usr/bin/env python3
#  -*- coding: utf-8 -*-
#
# yumizi @ 2023-06
#
# 一堆屎山代码
#
import json
import time
from threading import Thread
from tkinter.filedialog import asksaveasfile

import docx

import log
import os
import tkinter as tk
from tkinter import filedialog
from tkinter.messagebox import askyesno, showinfo, showwarning

from gui import ApplicationGUI, ExportGUI

# 版本
from oletools import oleobj

v = 'v 0.1.0.1'

# 日志对象
logger = log.Logger().log_create()

# APP根目录
app_path = os.path.dirname(os.path.abspath(__file__))


def file_check(file_path: str):
    if os.path.basename(file_path).startswith('~$'):
        logger.debug('未添加：缓存文件 %s', file_path)
        return False
    if not os.path.basename(file_path).endswith('.docx'):
        logger.debug('未添加：非docx文件 %s', file_path)
        return False
    if not os.path.isfile(file_path):
        logger.debug('未添加：不是文件 %s', file_path)
        return False
    return True


class Application(ApplicationGUI):
    def __init__(self, master=None, version=''):
        super().__init__(master)
        # 导出列表
        self.export_list = {}
        # 失败列表
        self.fail_list = []

        self.master = master
        self.version = version

        # 待处理文件列表
        self.file_list = []
        # 文件添加索引
        self.file_index = 0

        self.main_listbox = None
        self.create_list(self.frame_list)
        self.set_combobox_text()
        self.set_checkbutton_text()
        self.register_command()

        self.entry_tips_val.set('<< (>_<)')

        self.export_ui = None
        self.e_progress = 0
        self.e_total_task = 0
        self.e_rename = 0
        self.e_cancel = 0
        self.e_pause = 0

        self.lb_brand.configure(text='批量导出docx文本、图片和附件程序\n' + self.version + '\n@B站：敏Ymm')

    def set_checkbutton_text(self):
        self.che_text.set(1)
        self.che_image.set(1)
        self.che_table.set(1)
        self.che_combine.set(1)
        self.che_attachment.set(1)
        self.che_info.set(1)

    def set_combobox_text(self):
        self.cb_save_position['value'] = ('原文件目录中以原文件命名的子文件夹中', '原文件所在文件夹', '自定义文件夹')
        self.combobox_save_path.set(self.cb_save_position['value'][0])
        self.cb_save_position.current(0)

        self.cb_export_cover['value'] = ('同路径文件：覆盖', '同路径文件：自动重命名', '同路径文件：跳过')
        self.combobox_export_cover.set(self.cb_export_cover['value'][0])

        self.cb_name_part1['value'] = ('|自增编号|', '|连接符|', '|原文件名|', '|后缀名|')
        self.cb_name_part2['value'] = ('|连接符|', '|原文件名|', '|后缀名|', '|自增编号|')
        self.cb_name_part3['value'] = ('|原文件名|', '|后缀名|', '|自增编号|', '|连接符|')
        self.cb_name_part4['value'] = ('|后缀名|', '|自增编号|', '|连接符|', '|原文件名|')

        self.combobox_name1.set(self.cb_name_part1['value'][0])
        self.combobox_name2.set(self.cb_name_part2['value'][0])
        self.combobox_name3.set(self.cb_name_part3['value'][0])
        self.combobox_name4.set(self.cb_name_part4['value'][0])
        self.cb_name_part1.current(0)
        self.cb_name_part2.current(0)
        self.cb_name_part3.current(0)
        self.cb_name_part4.current(0)

    def create_list(self, frame):
        # 一个列表
        list_frame = tk.Frame(frame)
        scroll_h_bar = tk.Scrollbar(list_frame, orient=tk.HORIZONTAL)  # 水平滚动条组件
        scroll_v_bar = tk.Scrollbar(list_frame, orient=tk.VERTICAL, )  # 垂直滚动条组件
        self.main_listbox = tk.Listbox(list_frame,
                                       width=60, height=20,
                                       selectmode=tk.MULTIPLE,
                                       yscrollcommand=scroll_v_bar.set,
                                       xscrollcommand=scroll_h_bar.set)

        scroll_v_bar.pack(side=tk.RIGHT, fill=tk.Y)  # 设置垂直滚动条显示的位置
        scroll_v_bar.config(command=self.main_listbox.yview)  # 设置Scrollbar组件的command选项为该组件的yview()方法
        scroll_h_bar.pack(side=tk.BOTTOM, fill=tk.X)  # 设置水平滚动条显示的位置
        scroll_h_bar.config(command=self.main_listbox.xview)  # 设置Scrollbar组件的command选项为该组件的xview()方法
        self.main_listbox.place(relx=0, rely=0, relheight=0.95, relwidth=0.952, bordermode='ignore')
        list_frame.place(relx=0, rely=0, relheight=1, relwidth=1, bordermode='ignore')

    def register_command(self):
        # 选择文件
        self.btn_import_files.bind('<Button>', self.choose_file)
        # 选择文件夹
        self.btn_import_dir.bind('<Button>', self.choose_dir)
        # 删除列表选中项
        self.btn_delete_list_items.bind('<Button>', self.remove_list_item)
        # 删除所有项
        self.btn_delete_list_all.bind('<Button>', self.remove_list_all)
        # 选择文件夹
        self.btn_choose_position.bind('<Button>', self.choose_export_dir)
        # 导出按钮
        self.btn_export.bind('<Button-1>', self.export)
        # ---------------------------------------------------------------------------
        self.bind_cb_evt(self.cb_name_part1, self.name_tips)
        self.bind_cb_evt(self.cb_name_part2, self.name_tips)
        self.bind_cb_evt(self.cb_name_part3, self.name_tips)
        self.bind_cb_evt(self.cb_name_part4, self.name_tips)
        self.bind_cb_evt(self.cb_export_cover, self.export_tips_cover)
        # ---------------------------------------------------------------------------
        self.bind_cb_evt(self.cb_save_position, self.dir_tips)
        self.bind_cb_evt(self.entry_save_position, self.dir_tips)
        # ---------------------------------------------------------------------------
        self.cb_delete_raw_file.bind('<Button>', self.delete_tips)
        # ---------------------------------------------------------------------------
        self.cb_export_type_text.bind('<Button>', self.export_tips_text)
        self.cb_export_type_image.bind('<Button>', self.export_tips_image)
        self.cb_export_type_attachment.bind('<Button>', self.export_tips_attachment)
        self.cb_export_type_table.bind('<Button>', self.export_tips_table)
        self.cb_export_type_combine_text_table.bind('<Button>', self.export_tips_combine_text_table)
        self.cb_export_type_info.bind('<Button>', self.export_tips_info)

    @staticmethod
    def bind_cb_evt(cb, evt):
        cb.bind('<Button>', evt)
        cb.bind('<<ComboboxSelected>>', evt)
        cb.bind('<space>', evt)
        cb.bind('<Return>', evt)
        cb.bind('<Key>', evt)

    def export_tips_cover(self, evt):
        if not evt:
            return
        way = self.combobox_export_cover.get()
        if way == self.cb_export_cover['value'][0]:
            # 覆盖
            self.entry_tips_val.set('导出时相同路径的文件会被覆盖掉')
        elif way == self.cb_export_cover['value'][1]:
            # 重命名
            self.entry_tips_val.set('导出时若文件已存在则自动重命名')
        elif way == self.cb_export_cover['value'][2]:
            # 跳过
            self.entry_tips_val.set('导出时若文件已存在则跳过')
        else:
            # 意外
            self.entry_tips_val.set('这是什么情况？？！')

    def export_tips_image(self, evt):
        if not evt:
            return
        # 这个值是点击之前的值，
        if not self.che_image.get():
            self.entry_tips_val.set('导出Word中的图片')
        else:
            self.entry_tips_val.set('不导出Word中的图片')

    def export_tips_attachment(self, evt):
        if not evt:
            return
        # 这个值是点击之前的值，
        if not self.che_attachment.get():
            self.entry_tips_val.set('导出Word中的附件')
        else:
            self.entry_tips_val.set('不导出Word中的附件')

    def export_tips_table(self, evt):
        if not evt:
            return
        # 这个值是点击之前的值，
        if not self.che_table.get():
            self.entry_tips_val.set('导出Word表格中的文字')
        else:
            self.entry_tips_val.set('不导出Word表格中的文字')

    def export_tips_info(self, evt):
        if not evt:
            return
        # 这个值是点击之前的值，
        if not self.che_info.get():
            self.entry_tips_val.set('导出Word文档的信息（作者、修改时间啥的）')
        else:
            self.entry_tips_val.set('不导出Word文档的信息')

    def export_tips_combine_text_table(self, evt):
        if not evt:
            return
        # 这个值是点击之前的值，
        if not self.che_combine.get():
            self.entry_tips_val.set('合并导出Word中的普通文字和表格文字')
        else:
            self.entry_tips_val.set('分别导出Word中的普通文字和表格文字')

    def export_tips_text(self, evt):
        if not evt:
            return
        # 这个值是点击之前的值，
        if not self.che_text.get():
            self.entry_tips_val.set('导出Word中的普通文字')
        else:
            self.entry_tips_val.set('不导出Word中的普通文字')

    def delete_tips(self, evt):
        if not evt:
            return
        # 这个值是点击之前的值，
        if not self.che_delete_raw.get():
            self.entry_tips_val.set('导出成功后立即删除原文件')
        else:
            self.entry_tips_val.set('导出成功后保留原文件')

    def dir_tips(self, evt):
        if not evt:
            return
        way = self.combobox_save_path.get()
        if way == self.cb_save_position['value'][0]:
            # 原文件子目录
            self.entry_tips_val.set('会在原文件所在目录创建一个同名子文件夹以存储')
        elif way == self.cb_save_position['value'][1]:
            # 原文件目录
            self.entry_tips_val.set('保存位置与原文件在同一文件夹')
        else:
            export_dir = self.entry_save_position_val.get()
            # 指定目录
            self.entry_tips_val.set('保存位置：' + export_dir)

    def name_tips(self, evt):
        if not evt:
            return
        name = self.combobox_name1.get()
        name += self.combobox_name2.get()
        name += self.combobox_name3.get()
        name += self.combobox_name4.get()

        name = name.replace('|自增编号|', '编号') \
            .replace('|连接符|', ' - ') \
            .replace('|原文件名|', '原文件名') \
            .replace('|后缀名|', '.后缀') \
            .strip()

        self.entry_tips_val.set('文件名格式：%s' % name)

    def choose_export_dir(self, evt):
        if not evt:
            return
        directory = filedialog.askdirectory()
        if directory:
            self.entry_save_position_val.set(directory)
        self.dir_tips(None)

    def remove_list_all(self, evt):
        if not evt:
            return
        if not askyesno('删除所有项', '是否清空已导入列表？'):
            return

        if isinstance(self.main_listbox, tk.Listbox):
            self.main_listbox.delete(0, 'end')

        # 待处理文件列表
        self.file_list = []
        # 文件添加索引
        self.file_index = 0

        self.entry_tips_val.set('已清空列表')

    def remove_list_item(self, evt):
        if not evt:
            return
        if len(self.main_listbox.curselection()) == 0:
            showwarning('错误', '没有选中任何条目')
            self.entry_tips_val.set('没有从列表中移除任何文件')
        else:
            success = 0
            remove_index_list = []
            for index in self.main_listbox.curselection():
                remove_index_list.append(index)
                self.main_listbox.select_clear(index)

            # 需要从后面开始删除
            remove_index_list.sort(reverse=True)

            remove_file_list = []
            for index in remove_index_list:
                item_text = self.main_listbox.get(index)
                if item_text:
                    file_path = str(item_text).split('|-', 1)[1]
                    remove_file_list.append(file_path)
                    self.main_listbox.delete(index)

            for file_path in remove_file_list:
                self.file_list.remove(file_path)
                success += 1

            self.entry_tips_val.set('已从列表移除{0}个文件'.format(success))

    def add_file_list(self, files, is_choose_son=False):
        file_list = []

        # 遍历目录，拿到文件列表
        if isinstance(files, str) and os.path.isdir(files):
            logger.debug('添加文件夹')
            directory = files
            logger.debug('添加子文件夹中的内容：%s 选择目录：%s', is_choose_son, directory)
            for root, dirs, files in os.walk(directory):  # 遍历目录
                if directory != root and not is_choose_son:  # 跳过子文件夹
                    continue
                logger.debug('添加目录：%s', root)
                for file in files:  # 遍历文件
                    file_path = os.path.join(root, file)  # 拼接路径
                    file_list.append(file_path)  # 添加到列表中
        elif isinstance(files, tuple):
            logger.debug('添加文件')
            file_list = files
        else:
            showwarning('提示', '导入文件参数不正确')

        success = 0

        # 序号前面补 0
        bit = len(str(len(file_list)))
        if bit < 4:
            bit = 4

        # 添加到列表视图中
        for file in file_list:
            if file_check(file):
                # 已添加，不需要再添加
                if file in self.file_list:
                    continue
                self.file_index += 1
                logger.debug('添加文件：%s', file)
                self.main_listbox.insert(0, str(self.file_index).rjust(bit, '0') + '|-' + file)  # 从最后一个位置开始加入值
                self.file_list.append(file)
                success += 1

        tip_str = '成功添加{0}个docx文件，失败{1}个'.format(success, len(file_list) - success)
        self.entry_tips_val.set(tip_str)
        logger.debug(tip_str)
        showinfo('添加结果', '成功添加{0}个docx文件'.format(success))

    def choose_file(self, evt):
        if not evt:
            return
        self.entry_tips_val.set('正在导入，请等待.........')
        files_tuple = filedialog \
            .askopenfilename(title='请选择docx文件', filetypes=[('Word', '.docx')],
                             defaultextension='.docx',
                             multiple=True)
        if files_tuple:
            Thread(target=self.add_file_list, args=(files_tuple,)).start()

    def choose_dir(self, evt):
        """
        选择文件夹
        :param evt: 事件
        :return: None
        """
        if not evt:
            return
        self.entry_tips_val.set('正在导入，请等待.........')
        is_choose_son = askyesno('选择文件夹', '选择文件夹时是否选择子文件夹内的文件？')
        directory = filedialog.askdirectory()
        if directory:
            Thread(target=self.add_file_list, args=(directory, is_choose_son)).start()

    def export(self, evt):
        """
        导出按钮事件，负责校验参数，整合参数
        :param evt: 事件
        :return: None
        """
        if not evt:
            return
        export_dir_choose = self.combobox_save_path.get()
        logger.debug('导出位置：%s', export_dir_choose)

        export_dir = self.entry_save_position_val.get()
        logger.debug('导出目录：%s', export_dir)

        if export_dir_choose == self.cb_save_position['value'][2]:
            if not (export_dir and os.path.isdir(export_dir)):
                showwarning('导出时遇到问题', '保存位置未设置文件夹，请设置保存文件夹或设置其他保存位置')
                return

        try:
            export_dir_choose = self.cb_save_position['value'].index(export_dir_choose) + 1
            logger.debug('导出方式：%s', export_dir_choose)
        except ValueError:
            showwarning('出错', '保存方式参数错误')
            return

        name = self.combobox_name1.get()
        name += self.combobox_name2.get()
        name += self.combobox_name3.get()
        name += self.combobox_name4.get()
        logger.debug('名字规则：%s', name)
        if not name.endswith('|后缀名|'):
            if not askyesno('提示', '文件名最好以“|后缀名|”结尾，否则可能导致无法识别，是否继续？'):
                return

        export_type = []
        if self.che_text.get():
            export_type.append('text')
        if self.che_table.get():
            export_type.append('table')
        if self.che_image.get():
            export_type.append('image')
        if self.che_attachment.get():
            export_type.append('attachment')
        if self.che_combine.get():
            export_type.append('combine')
        if self.che_info.get():
            export_type.append('info')

        logger.debug('导出类型：%s', str(export_type))

        is_delete = self.che_delete_raw.get()
        logger.debug('导出后删除原文件：%s', str(is_delete))

        cover = self.combobox_export_cover.get()
        if cover == self.cb_export_cover['value'][0]:
            # 覆盖
            cover = 'cover'
            logger.debug('导出时相同路径的文件会被覆盖掉')
        elif cover == self.cb_export_cover['value'][1]:
            # 重命名
            cover = 'rename'
            logger.debug('导出时若文件已存在则自动重命名')
        elif cover == self.cb_export_cover['value'][2]:
            # 跳过
            cover = 'skip'
            logger.debug('导出时若文件已存在则跳过')
        else:
            # 意外
            showwarning('提示', '参数出现了意外，未知的文件覆盖策略')

        parameter = {
            'way': export_dir_choose,  # 保存方式
            'dir': export_dir,  # 保存目录
            'name': name,  # 文件名格式
            'type': export_type,  # 导出类型
            'del': bool(is_delete),  # 导出后删除原文件
            'list': self.file_list,  # 导出文件
            'cover': cover,  # 覆盖策略
        }

        # 开始导出
        self.start_export(parameter)

    # ====================================================================================

    def start_export(self, parameter):
        """
        初始化导出界面，启动导出线程
        :param parameter: 导出参数
        :return: None
        """
        logger.debug('启动界面')
        # 初始化导出界面参数
        self.e_cancel = 0
        self.e_pause = 0
        self.e_progress = 0
        self.e_total_task = len(parameter.get('list'))
        # 初始化界面
        self.export_ui = ExportGUI(self.master)
        # 中间文本框初始化
        self.export_ui.txt_d_result_show.delete('1.0', 'end')
        self.export_ui.txt_d_result_show.insert('end', '准备导出......\n')
        # 切换到进度条界面
        self.export_ui.switch_func(None, '暂停', '取消', None, None)
        # 重置提示信息
        self.export_ui.lb_d_tips_var.set('''处理进度：∞%''')
        # 重置进度和选项
        self.export_ui.pd_d_main_var.set(0)  # 进度
        self.export_ui.cb_d_all_var.set(0)  # 一律如此选项
        # 绑定点击事件
        self.export_ui.btn_d_e.bind('<Button-1>', self.e_btn_ex)
        self.export_ui.btn_d_left.bind('<Button-1>', self.e_btn_left)
        self.export_ui.btn_d_right.bind('<Button-1>', self.e_btn_right)

        Thread(target=self.run, args=(parameter,)).start()

    def show_fail(self):
        for i in range(7):
            time.sleep(0.05)
            self.export_ui.txt_d_result_show.insert(1.0, '-> \n')
        for i, file in enumerate(self.fail_list):
            time.sleep(0.01)
            self.export_ui.txt_d_result_show.insert(1.0, '失败{0} -> {1}\n'.format(str(i + 1).rjust(3, '0'), file))
        show_msg = """
-->> 失败的原因可能是文件损坏或无内容，请尝试用WPS打开并另存为 <<--
-->> >>>>>>复制到文件管理器地址栏中Enter可直接打开<<<<<<< <<--
\n"""
        show_msg = list(show_msg)
        show_msg.reverse()
        for i, c in enumerate(show_msg):
            if str(c) == '\n':
                time.sleep(0.4)
            else:
                time.sleep(0.015)
            self.export_ui.txt_d_result_show.insert(1.0, str(c))

    def e_btn_ex(self, evt):
        """
        点击“扩展”按钮
        :param evt: 事件
        :return: None
        """
        if not evt:
            return
        if isinstance(self.export_ui, ExportGUI):
            btn_text = self.export_ui.btn_d_e_var.get()
            if btn_text == '失败列表':
                Thread(target=self.show_fail).start()

    def e_btn_left(self, evt):
        """
        点击“暂停”按钮
        :param evt: 事件
        :return: None
        """
        if not evt:
            return
        if isinstance(self.export_ui, ExportGUI):
            btn_text = self.export_ui.btn_d_left_var.get()
            if btn_text == '暂停':
                self.e_pause = 1
                self.export_ui.btn_d_left_var.set('继续')
                self.show_progress(0, '.............................', '已暂停')
            elif btn_text == '继续':
                self.e_pause = 0
                self.export_ui.btn_d_left_var.set('暂停')
            elif btn_text == '导出报告':
                self.save_json(self.export_list)

    def e_btn_right(self, evt):
        """
        点击“取消”按钮
        :param evt: 事件
        :return: None
        """
        if not evt:
            return
        if isinstance(self.export_ui, ExportGUI):
            btn_text = self.export_ui.btn_d_right_var.get()
            if btn_text == '取消':
                self.e_cancel = 1
            elif btn_text == '关闭':
                self.export_ui.tf_d_title.place(relx=0.0, rely=0.0, relheight=0.0, relwidth=0.0)
                self.export_ui.tf_d_title.destroy()

    def show_progress(self, current, file, state):
        """
        显示进度
        :param current: 当前序号
        :param file: 文件
        :param state: 导出状态
        :return: None
        """

        # 计算百分比
        if self.e_total_task != 0:
            current = (current + 1) / self.e_total_task * 100
        else:
            current = 0

        # 百分比不能掉（单线程好像没必要）
        if self.e_progress > current:
            current = self.e_progress
        else:
            self.e_progress = current

        # 设置界面显示百分比
        self.export_ui.pd_d_main_var.set(current)
        self.export_ui.lb_d_tips_var.set('处理进度：{0}%'.format(current))
        # 在文本框中显示细节
        self.export_ui.txt_d_result_show.insert(1.0, '{0} -> {1}\n'.format(state, file))

    def run(self, parameter):
        print('开始处理任务')

        success = 0
        sub_file_count = 0

        # 处理参数
        file_list = parameter.get('list')
        if not isinstance(file_list, list):
            print('参数错误')
            return

        # 清空上次报告
        self.export_list = {}

        self.fail_list = []

        for index, file in enumerate(file_list):
            self.export_list[file] = []
            time.sleep(0.01)
            if self.e_cancel:
                if askyesno('取消导出', '是否取消导出？（已导出的不会被清理）'):
                    break
            # 暂停
            while self.e_pause:
                continue
            logger.debug('处理文件：%s', file)
            try:
                sub_file_list = self.dispose(index, file, parameter)
                if isinstance(sub_file_list, bool) and not sub_file_list:
                    raise IOError('文件打开出错')
                if not sub_file_list:
                    sub_file_list = []
            except Exception as e:
                logger.exception('导出错误：导出文件时出错 %s' % str(e))
                self.fail_list.append(file)
                self.show_progress(index, file, '失败\t'.rjust(6, '-'))
                continue
            logger.debug('成功 %s', file)
            # 统计计数
            sub_count = len(sub_file_list)
            success += 1
            sub_file_count += sub_count
            # 保存到总的导出列表中
            self.export_list[file] = sub_file_list
            self.show_progress(index, file, '成功 %s\t' % str(sub_count).rjust(2, '0'))

        if len(file_list) == 0:
            logger.debug('没有需要处理的文件')
            self.entry_tips_val.set('没有需要处理的文件')
            self.export_ui.txt_d_result_show.insert(1.0, '没有需要处理的文件\n')
            self.export_ui.switch_func(None, None, '关闭', None, None)
            showinfo('提示', '没有需要处理的文件')
        else:
            # 总结提示
            showinfo('导出结果',
                     '导出完成\n成功导出%d个docx文档\n生成子文档%d个\n失败%d个' % (success, sub_file_count, len(file_list) - success))

            self.entry_tips_val.set('导出完成')
            logger.debug('任务完成')
            if self.fail_list:
                self.export_ui.switch_func('失败列表', '导出报告', '关闭', None, None)
            else:
                self.export_ui.switch_func(None, '导出报告', '关闭', None, None)
            # self.save_json(export_list)

    @staticmethod
    def save_json(export_list):
        f = asksaveasfile(mode='wb', defaultextension=".json")
        if f:
            json_str = json.dumps(export_list, ensure_ascii=False)
            f.write(json_str.encode(encoding='utf-8'))
            f.close()
            logger.debug('导出结果保存成功')
            showinfo('导出', '导出成功！')
        else:
            logger.debug('用户取消保存导出结果')

    @staticmethod
    def get_new_path(i, seek, file_path, file_name_format, out_dir, cover='rename'):
        """
        获取一个新的文件路径
        :param cover: 覆盖策略 rename重命名（默认），skip跳过，cover覆盖
        :param i: 文件编号
        :param seek: 导出附件编号
        :param file_path: 文件名（路径）
        :param out_dir: 输出目录
        :param file_name_format: 文件名格式
        :return: 新的文件路径
        """
        if not os.path.isdir(out_dir):
            out_dir = os.path.dirname(os.path.abspath(out_dir))  # 获取导出路径
        if file_name_format is None:
            file_name_format = '|自增编号||连接符||原文件名||后缀名|'
        number = str(i).rjust(4, '0') + '.' + str(seek).rjust(2, '0')  # |自增编号|，不够前面添0
        raw_file_name = os.path.basename(file_path)  # |原文件名|
        suffix = ''  # |后缀名|
        if '.' in raw_file_name:
            name_part = raw_file_name.split('.')
            part_count = len(name_part)
            suffix = name_part[-1]  # 后缀名
            raw_file_name = '.'.join(name_part[:(part_count - 1)])  # 前面的，这么处理是因为有的文件名中有多个 .
        link_char = ' - '  # |连接符|
        file_name = out_dir + '/' + file_name_format.replace('|自增编号|', number) \
            .replace('|连接符|', link_char) \
            .replace('|原文件名|', raw_file_name) \
            .replace('|后缀名|', '.' + suffix) \
            .strip()
        if 'cover' == cover:
            return file_name  # 拼接路径
        if 'skip' == cover:
            if os.path.isfile(file_name):
                return None
            else:
                return file_name
        if 'rename' == cover:
            add_index = 1
            while True:
                if not os.path.isfile(file_name):
                    return file_name
                file_name = out_dir + '/' + file_name_format.replace('|自增编号|', number) \
                    .replace('|连接符|', link_char) \
                    .replace('|原文件名|', raw_file_name) \
                    .replace('|后缀名|', '.' + str(add_index) + '.' + suffix) \
                    .strip()
                add_index += 1

    @staticmethod
    def re_decode(s: str, encoding: str = 'gbk'):
        """
        重新解码，解决oleobj对中文乱码的问题
        :param s: 原始字符串
        :param encoding: 新的解码编码，默认为 GBK
        :return: 新的字符串
        """
        i81 = s.encode('iso-8859-1')
        return i81.decode(encoding)

    def dispose(self, index: int, docx_file: str, parameter):
        """
        处理docx文档
        :param index: 索引，文档级
        :param docx_file: 文档路径
        :param parameter: 其他参数
        :return: 文档导出列表
        """
        save_way = parameter.get('way')
        output_dir = parameter.get('dir')
        name_format = parameter.get('name')
        export_type = parameter.get('type')
        is_del_raw = parameter.get('del')
        cover = parameter.get('cover')

        # 导出过程中是否出错
        is_error = False

        # 导出的文件列表
        export_files = []
        # 从1开始
        index += 1

        # 附件编号
        seek = 0

        if not docx_file.endswith('.docx'):
            logger.debug('不支持的文件类型')
            return

            # 导出方式设置
        if save_way == 1:
            output_dir = docx_file[:-5]
        elif save_way == 2:
            output_dir = os.path.dirname(docx_file)
        else:
            if output_dir:
                if not os.path.isdir(output_dir):
                    os.makedirs(output_dir)
            else:
                logger.error('导出路径不正确：' + docx_file)
                return export_files

        # 导出文件前后取出空格（不去除可能导出失败）
        output_dir = output_dir.strip()

        # 创建导出文件夹
        if not os.path.isdir(output_dir):
            logger.debug('创建导出文件夹：%s', output_dir)
            os.mkdir(output_dir)
            if os.path.isdir(output_dir):
                logger.debug('创建导出文件夹成功')
            else:
                logger.debug('创建导出文件夹失败')

        # 打开docx文档
        try:
            docx_document = docx.Document(docx_file)
        except Exception as e:
            logger.exception('打开文档时出错%s' % str(e))
            self.remove_empty_dir(output_dir)
            return False
        logger.debug('打开文档完成')

        # 文档信息
        if 'info' in export_type:
            docx_properties = docx_document.core_properties
            all_properties = '作者\t' + str(docx_properties.author) + '\n'
            all_properties += '类别\t' + str(docx_properties.category) + '\n'
            all_properties += '注释\t' + str(docx_properties.comments) + '\n'
            all_properties += '内容状态\t' + str(docx_properties.content_status) + '\n'
            all_properties += '创建时间\t' + str(docx_properties.created) + '\n'
            all_properties += '标识符\t' + str(docx_properties.identifier) + '\n'
            all_properties += '关键字\t' + str(docx_properties.keywords) + '\n'
            all_properties += '语言\t' + str(docx_properties.language) + '\n'
            all_properties += '最后修改者\t' + str(docx_properties.last_modified_by) + '\n'
            all_properties += '上次打印\t' + str(docx_properties.last_printed) + '\n'
            all_properties += '修改时间\t' + str(docx_properties.modified) + '\n'
            all_properties += '修订\t' + str(docx_properties.revision) + '\n'
            all_properties += '主题\t' + str(docx_properties.subject) + '\n'
            all_properties += '标题\t' + str(docx_properties.title) + '\n'
            all_properties += '版本\t' + str(docx_properties.version) + '\n'
            logger.debug('文档信息：%s', all_properties.replace('\n', '， '))

            # 导出文档信息
            seek += 1
            info_file_path = self.get_new_path(index, seek, '文档信息.txt', name_format, output_dir, cover)
            if info_file_path:
                logger.debug('文档信息保存位置：%s', info_file_path)
                with open(info_file_path, 'w', encoding='utf-8') as f:
                    f.write(all_properties)
                export_files.append(info_file_path)
            else:
                logger.debug('跳过文件')

        # 所有文本
        all_text = ''
        if 'text' in export_type:
            for paragraph in docx_document.paragraphs:
                all_text += paragraph.text + ' '  # 段落之间用空格隔开
            logger.debug('所有文本：%s', all_text)

        # 所有表格
        all_table_text = ''
        if 'table' in export_type:
            for table in docx_document.tables:
                for cell in getattr(table, '_cells'):
                    all_table_text += cell.text + '|'  # 单元格之间用 “|” 隔开
            logger.debug('所有表格文本：%s', all_table_text)

        # 导出文本
        if 'combine' in export_type and ('text' in export_type or 'table' in export_type):

            seek += 1
            combine_file_path = self.get_new_path(index, seek, '文本和表格.txt', name_format, output_dir,
                                                  cover)
            if combine_file_path:
                logger.debug('文档文本和表格保存位置：%s', combine_file_path)
                with open(combine_file_path, 'w', encoding='utf-8') as f:
                    f.write(all_text)
                    f.write('\n')  # 换个行
                    f.write(all_table_text)
                export_files.append(combine_file_path)
            else:
                logger.debug('跳过文件')
        else:
            if 'text' in export_type:
                seek += 1
                text_file_path = self.get_new_path(index, seek, '文本.txt', name_format, output_dir, cover)
                if text_file_path:
                    logger.debug('文档文本保存位置：%s', text_file_path)
                    with open(text_file_path, 'w', encoding='utf-8') as f:
                        f.write(all_text)
                    export_files.append(text_file_path)
                else:
                    logger.debug('跳过文件')

            if 'table' in export_type:
                seek += 1
                table_file_path = self.get_new_path(index, seek, '表格.txt', name_format, output_dir, cover)
                if table_file_path:
                    logger.debug('文档表格保存位置：%s', table_file_path)
                    with open(table_file_path, 'w', encoding='utf-8') as f:
                        f.write(all_table_text)
                    export_files.append(table_file_path)
                else:
                    logger.debug('跳过文件')

        # 遍历所有附件
        if 'image' in export_type or 'attachment' in export_type:
            docx_related_parts = docx_document.part.related_parts
            for part in docx_related_parts:
                part = docx_related_parts[part]
                part_name = str(part.partname)  # 附件路径（partname）

                # 只导出这两个目录下的
                if not (part_name.startswith('/word/media/') or part_name.startswith('/word/embeddings/')):
                    continue

                # 构建导出路径
                seek += 1
                save_path = self.get_new_path(index, seek, part.partname, name_format, output_dir, cover)

                # ole 文件判断
                # 不符合 .bin 作为后缀且文件名中有ole，则不被认为是OLE文件
                if not (part_name.lower().endswith('.bin') and 'ole' in part_name.lower()):
                    # 如果没有支持图片导出
                    if 'image' not in export_type:
                        continue

                    if save_path is None:
                        logger.debug('跳过文件')
                        continue

                    # 直接写入文件
                    logger.debug('图片导出路径：%s', save_path)
                    with open(save_path, 'wb') as f:
                        f.write(part.blob)
                    export_files.append(save_path)  # 记录文件

                    continue

                # 如果没有支持附件导出
                if 'attachment' not in export_type:
                    continue

                # 将字节数组传递给oleobj处理
                for ole in oleobj.find_ole(save_path, part.blob):
                    if ole is None:  # 没有找到 OLE 文件，跳过
                        continue

                    for path_parts in ole.listdir():  # 遍历OLE中的文件

                        # 判断是不是[1]Ole10Native，使用列表推导式忽略大小写，不是的话就不要继续了
                        if '\x01ole10native'.casefold() not in [path_part.casefold() for path_part in
                                                                path_parts]:
                            continue

                        stream = None
                        try:
                            # 使用 Ole File 打开 OLE 文件
                            stream = ole.openstream(path_parts)
                            opkg = oleobj.OleNativeStream(stream)
                        except IOError:
                            logger.debug('不是OLE文件：%s', path_parts)
                            if stream is not None:  # 关闭文件流
                                stream.close()
                            continue

                        # 打印信息
                        if opkg.is_link:
                            logger.debug('是链接而不是文件，跳过')
                            continue

                        ole_filename = self.re_decode(opkg.filename)
                        ole_src_path = self.re_decode(opkg.src_path)
                        ole_temp_path = self.re_decode(opkg.temp_path)
                        logger.debug('文件名：%s，原路径：%s，缓存路径：%s', ole_filename, ole_src_path, ole_temp_path)

                        # 生成新的文件名
                        seek += 1
                        filename = self.get_new_path(index, seek, ole_filename, name_format, output_dir, cover)

                        logger.debug('OLE附件导出路径：%s', filename)

                        if filename is None:
                            logger.debug('跳过')
                            continue

                        # 转存
                        try:
                            with open(filename, 'wb') as writer:
                                n_dumped = 0
                                next_size = min(oleobj.DUMP_CHUNK_SIZE, opkg.actual_size)
                                while next_size:
                                    data = stream.read(next_size)
                                    writer.write(data)
                                    n_dumped += len(data)
                                    if len(data) != next_size:
                                        logger.warning('想要读取 %d, 实际取得 %d', next_size, len(data))
                                        break
                                    next_size = min(oleobj.DUMP_CHUNK_SIZE, opkg.actual_size - n_dumped)
                            export_files.append(filename)  # 记录导出的文件
                        except Exception as exc:
                            is_error = True
                            logger.exception('在转存时出现错误', exc)
                        finally:
                            stream.close()
        logger.debug('导出的所有文件：%s', export_files)
        if not is_error and is_del_raw:
            logger.debug('删除原文件：%s', docx_file)
            os.remove(docx_file)
        self.remove_empty_dir(output_dir)
        return export_files

    @staticmethod
    def remove_empty_dir(target_dir):
        if os.path.isdir(target_dir):
            if not os.listdir(target_dir):
                logger.debug("删除空文件夹：%s", target_dir)
                os.removedirs(target_dir)


def run():
    root = tk.Tk()
    Application(master=root, version=v)

    def close_window():
        ans = askyesno(title='提示', message='是否关闭窗口？')
        if ans:
            root.destroy()
        else:
            return

    root.protocol('WM_DELETE_WINDOW', close_window)

    root.title('导出docx')
    root.iconbitmap('images/icon.ico')
    root.mainloop()


if __name__ == '__main__':
    run()
